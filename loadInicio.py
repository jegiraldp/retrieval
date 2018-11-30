from docx import Document
global documento
documento="acta3.docx"

def getDocumento():
    return str(documento)

def stopwords():
    f = open('archivos/stopwords.txt','r')
    txt = f.read()
    words = []
    for i in txt:
        words = txt.split("\n")
    f.close()
    return words

def leerDocumento():
    docu=Document('archivos/'+documento)
    texto = ""
    words=[]
    for i in docu.paragraphs:
        texto += i.text.lower() + " "

    texto=especiales(texto)
    texto+=leerTextoTablas(docu)
    words=texto.split()
    return words

def leerTextoTablas(docu):
    tablas = docu.tables

    impri = ""
    retornoFinal=""

    if (tablas):
        tablas = docu.tables
        for r in tablas[0].rows:
            for celda in r.cells:
                impri += " " + celda.text

        impri=impri.lower()
        impri=especiales(impri)
        retorno=impri.split()

        for r in retorno:
            retornoFinal+=" "+str(r)

    return retornoFinal

def especiales(texto):
    texto=texto.replace('(','')
    texto = texto.replace(')', '')
    texto = texto.replace(':', '')
    texto = texto.replace('+', '')
    texto = texto.replace('-', '')
    texto = texto.replace('*', '')
    texto = texto.replace('%', '')
    texto = texto.replace('#', '')
    texto = texto.replace('$', '')
    texto = texto.replace('&', '')
    texto = texto.replace('?', '')
    texto = texto.replace('¿', '')
    texto = texto.replace(';', '')
    texto = texto.replace(',', '')
    texto = texto.replace('~', '')
    texto = texto.replace('{','')
    texto = texto.replace('}','')
    texto = texto.replace('[','')
    texto = texto.replace(']','')
    texto = texto.replace('.','')
    return texto